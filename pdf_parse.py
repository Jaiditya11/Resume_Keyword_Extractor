# import os
# from PyPDF2 import PdfReader
# import re

# def extract_resume_texts(directory='./data'):
#     """
#     Extract text from all PDF resumes in the specified directory
#     Returns a list of texts, one for each resume
#     """
#     resume_texts = []
    
#     # Iterate through all files in the directory
#     for filename in os.listdir(directory):
#         if filename.endswith('.pdf'):
#             file_path = os.path.join(directory, filename)
            
#             try:
#                 # Create a PDF reader object
#                 reader = PdfReader(file_path)
                
#                 # Extract text from all pages
#                 text = ''
#                 for page in reader.pages:
#                     text += page.extract_text()
                
#                 resume_texts.append(text)
                
#             except Exception as e:
#                 print(f"Error processing {filename}: {str(e)}")
    
#     return resume_texts


# # Export the resume texts as a list
# resume_texts = extract_resume_texts()

#############################################################################
# import os
# from PyPDF2 import PdfReader
# from docx import Document

# def extract_resume_texts(directory='./data'):
#     """
#     Extract text from all PDF and Word resumes in the specified directory.
#     Returns a list of texts, one for each resume.
#     """
#     resume_texts = []
    
#     # Iterate through all files in the directory
#     for filename in os.listdir(directory):
#         file_path = os.path.join(directory, filename)
        
#         # Process PDF files
#         if filename.endswith('.pdf'):
#             try:
#                 reader = PdfReader(file_path)
#                 text = ''
#                 for page in reader.pages:
#                     text += page.extract_text()
#                 resume_texts.append(text)
#             except Exception as e:
#                 print(f"Error processing PDF {filename}: {str(e)}")
        
#         # Process Word (.docx) files
#         elif filename.endswith('.docx'):
#             try:
#                 doc = Document(file_path)
#                 text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
#                 resume_texts.append(text)
#             except Exception as e:
#                 print(f"Error processing Word file {filename}: {str(e)}")
    
#     return resume_texts

# # Usage example: Export the resume texts as a list
# resume_texts = extract_resume_texts(directory='./data')

# # Print the first few resumes for verification
# for i, resume in enumerate(resume_texts[:3], 1):
#     print(f"Resume {i}:\n{resume}\n{'-'*80}")

#######################################################################################################
# import os
# from PyPDF2 import PdfReader
# import re

# def extract_resume_texts(directory='./data'):
#     """
#     Extract text from all PDF resumes in the specified directory.
#     Returns a list of texts, one for each resume.
#     """
#     resume_texts = []
    
#     # Iterate through all files in the directory
#     for filename in os.listdir(directory):
#         if filename.endswith('.pdf'):
#             file_path = os.path.join(directory, filename)
            
#             try:
#                 # Create a PDF reader object
#                 reader = PdfReader(file_path)
                
#                 # Extract text from all pages
#                 text = ''
#                 for page in reader.pages:
#                     page_text = page.extract_text()
#                     if page_text:  # Avoid appending None
#                         text += page_text
                
#                 resume_texts.append(text)
                
#             except Exception as e:
#                 print(f"Error processing {filename}: {str(e)}")
#     print(resume_texts)
#     return resume_texts

# def chunk_resumes(resumes):
#     """
#     Splits multiple resumes into chunks by individual.
#     """
#     resume_chunks = []
#     for resume in resumes:
#         # Use the name or email as an identifier to split resumes
#         match = re.search(r"^(\w+ \w+)\n", resume)  # Look for the name at the start
#         if match:
#             name = match.group(1)  # Extract name
#             resume_chunks.append({"name": name, "resume": resume})
#     return resume_chunks

# # Extract resume texts from PDF files
# resume_texts = extract_resume_texts()

# # Process resumes and chunk them
# chunked_resumes = chunk_resumes(resume_texts)

# # Display chunked resumes
# for chunk in chunked_resumes:
#     print(f"Name: {chunk['name']}")
#     print(f"Resume: {chunk['resume'][:500]}")  # Display the first 500 characters for brevity
#     print("\n" + "-"*50 + "\n")


import os
from PyPDF2 import PdfReader
from pathlib import Path
from docx import Document

def extract_resume_texts(directory='./data'):
    """
    Extract text from all PDF and Word resumes in the specified directory.
    Returns a list of texts, one for each resume.
    """
    resume_texts = []
    
    # Iterate through all files in the directory
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        
        # Process PDF files
        if filename.endswith('.pdf'):
            try:
                reader = PdfReader(file_path)
                text = ''
                for page in reader.pages:
                    text += page.extract_text()
                resume_texts.append(text)
                print(f"Resume {filename}: {text}")
            except Exception as e:
                print(f"Error processing PDF {filename}: {str(e)}")
        
        # Process Word (.docx) files
        elif filename.endswith('.docx'):
            try:
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                resume_texts.append(text)
                print(f"Resume {filename}: {text}")
            except Exception as e:
                print(f"Error processing Word file {filename}: {str(e)}")
    
    return resume_texts
# Usage example: Export the resume texts as a list
resume_texts = extract_resume_texts(directory='./data')
# Print the first few resumes for verification
for i, resume in enumerate(resume_texts[:3], 1):
    print(f"Resume {i}: {resume}")